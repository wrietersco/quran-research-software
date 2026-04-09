"""Write Step 6 report artifacts: Markdown, HTML, DOCX, PDF (best-effort)."""

from __future__ import annotations

import re
import subprocess
from pathlib import Path


def write_report_markdown(markdown_body: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(markdown_body, encoding="utf-8")


def write_report_html(markdown_body: str, path: Path, *, title: str = "Report") -> None:
    import markdown

    path.parent.mkdir(parents=True, exist_ok=True)
    html_body = markdown.markdown(
        markdown_body,
        extensions=["fenced_code", "tables", "nl2br"],
        output_format="html",
    )
    doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1"/>
  <title>{_html_escape(title)}</title>
  <style>
    body {{ font-family: Georgia, serif; max-width: 48rem; margin: 2rem auto; line-height: 1.55; }}
    pre, code {{ font-family: ui-monospace, Consolas, monospace; font-size: 0.9em; }}
    pre {{ background: #f5f5f5; padding: 0.75rem; overflow-x: auto; }}
    h1, h2, h3 {{ font-family: system-ui, sans-serif; }}
    .arabic {{ direction: rtl; unicode-bidi: embed; font-size: 1.15em; }}
  </style>
</head>
<body>
{html_body}
</body>
</html>
"""
    path.write_text(doc, encoding="utf-8")


def _html_escape(s: str) -> str:
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def write_report_docx(markdown_body: str, path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    lines = markdown_body.splitlines()
    in_fence = False
    buf_fence: list[str] = []

    def flush_fence() -> None:
        nonlocal buf_fence
        if buf_fence:
            p = doc.add_paragraph("\n".join(buf_fence))
            p.style = "Intense Quote"
        buf_fence = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_fence:
                flush_fence()
                in_fence = False
            else:
                in_fence = True
            continue
        if in_fence:
            buf_fence.append(line)
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        elif not stripped:
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    doc.save(str(path))


def write_report_pdf(html_path: Path, pdf_path: Path) -> bool:
    """Try xhtml2pdf; then optional pandoc. Returns True if pdf_path was written."""
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    html = html_path.read_text(encoding="utf-8")
    try:
        from xhtml2pdf import pisa

        with pdf_path.open("wb") as fh:
            ok = pisa.CreateDocument(html, dest=fh).err == 0
        return bool(ok and pdf_path.is_file())
    except Exception:
        pass
    pandoc = shutil_which("pandoc")
    if pandoc:
        try:
            subprocess.run(
                [pandoc, str(html_path), "-o", str(pdf_path)],
                check=True,
                capture_output=True,
                text=True,
            )
            return pdf_path.is_file()
        except Exception:
            pass
    return False


def shutil_which(cmd: str) -> str | None:
    from shutil import which

    return which(cmd)


def write_all_formats(markdown_body: str, out_dir: Path, *, title: str = "Report") -> dict[str, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    md = out_dir / "report.md"
    html_p = out_dir / "report.html"
    docx_p = out_dir / "report.docx"
    pdf_p = out_dir / "report.pdf"
    write_report_markdown(markdown_body, md)
    write_report_html(markdown_body, html_p, title=title)
    write_report_docx(markdown_body, docx_p)
    ok_pdf = write_report_pdf(html_p, pdf_p)
    return {"md": md, "html": html_p, "docx": docx_p, "pdf": pdf_p if ok_pdf else None}
